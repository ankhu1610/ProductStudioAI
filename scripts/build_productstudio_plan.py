from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "ProductStudio_AI_Build_Plan.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
INK = "1E293B"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_font(run, size=11, bold=False, color=INK):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text, style=None, bold=False, color=INK):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_font(r, bold=bold, color=color)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        set_font(p.add_run(item))


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        set_font(p.add_run(item))


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_width(cell, widths[idx])
        set_cell_margins(cell)
        set_cell_shading(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(header), size=10, bold=True, color=DARK_BLUE)
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            if row_idx % 2:
                set_cell_shading(cell, LIGHT_GRAY)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            set_font(p.add_run(value), size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    set_font(p.add_run(text), size={1: 16, 2: 13, 3: 12}[level], bold=True,
             color=BLUE if level < 3 else DARK_BLUE)
    return p


def add_phase(doc, title, goal, tasks, success):
    heading(doc, title, 2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    set_font(p.add_run("Goal: "), bold=True, color=DARK_BLUE)
    set_font(p.add_run(goal))
    add_bullets(doc, tasks)
    p = doc.add_paragraph()
    set_font(p.add_run("Success criterion: "), bold=True, color=DARK_BLUE)
    set_font(p.add_run(success))


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for level, before, after in ((1, 18, 10), (2, 14, 7), (3, 10, 5)):
        style = styles[f"Heading {level}"]
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("PRODUCTSTUDIO AI")
    set_font(r, size=25, bold=True, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    set_font(p.add_run("Build Plan: Mask-Aware Latent Diffusion for E-Commerce Ad Creative"), size=13, color="475569")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    set_font(p.add_run("Prepared for local development on an RTX 5070 Ti Laptop GPU (12 GB VRAM)"), size=10, color="64748B")

    heading(doc, "Project Goal")
    add_text(doc, "Build ProductStudio AI: an e-commerce ad-creative engine that preserves a merchant's original product pixels while generating prompt-controlled studio or lifestyle backgrounds.")
    heading(doc, "Core User Flow", 2)
    add_numbered(doc, [
        "Merchant uploads a product image and a binary product mask.",
        "Merchant enters a background prompt, such as 'premium sneaker campaign on wet neon asphalt at night'.",
        "The inpainting pipeline generates the surrounding scene while preserving the masked product region.",
        "The service composites the original foreground back into the result, then saves the image and reproducibility metadata."
    ])

    heading(doc, "Scope Boundaries")
    add_table(doc, ["Build in MVP", "Defer until later"], [
        ("512x512 mask-guided background replacement", "Automatic product segmentation"),
        ("Stable Diffusion 1.5 inpainting", "Product-specific LoRA training"),
        ("DDIM sampling with 10-50 configurable steps", "Redis queue and multiple GPU workers"),
        ("FastAPI, Gradio UI, Docker, metadata logging", "Kubernetes, Grafana, cloud-scale deployment"),
        ("Local MLflow experiments and benchmark report", "SDXL and full-model fine-tuning")
    ], [4680, 4680])

    heading(doc, "System Constraints and Decisions")
    add_table(doc, ["Area", "Decision / Constraint"], [
        ("Hardware", "Core Ultra 9 275HX, 32 GB RAM, RTX 5070 Ti Laptop GPU with 12 GB VRAM, ample local storage."),
        ("Resolution", "Start and benchmark at 512x512 only."),
        ("GPU memory", "Use FP16/mixed precision; internal batch size 1; generate multiple variants sequentially."),
        ("Concurrency", "Run one GPU inference worker to prevent out-of-memory failures."),
        ("Model", "Use Stable Diffusion 1.5 inpainting for the product-facing application."),
        ("Product fidelity", "Use mask-aware generation plus final pixel-space compositing; do not claim latent blending alone is pixel-perfect."),
        ("Latency", "Measure p50/p95 latency and GPU memory on this laptop before making performance claims."),
        ("Training", "LoRA is optional phase 5. Full model fine-tuning is explicitly out of scope."),
        ("Licensing", "Use authorized product photos and confirm checkpoint/model license terms before sharing a demo.")
    ], [1900, 7460])

    heading(doc, "Target Architecture")
    add_text(doc, "Merchant / Gradio UI -> FastAPI API -> validation -> generation service -> SD 1.5 inpainting with DDIM -> pixel-space compositing -> image and JSON metadata. MLflow stores experiment artifacts; SQLite stores generation records.")
    add_bullets(doc, [
        "Keep the current from-scratch Stable Diffusion code as a documented research component that demonstrates architectural understanding.",
        "Use maintained Diffusers components for the product service because they are designed for reliable mask-conditioned inpainting.",
        "Begin with local filesystem storage; move to object storage only if a cloud deployment is added."
    ])

    heading(doc, "Development Roadmap")
    add_phase(doc, "Phase 0 - Foundation (2-3 days)", "Create a reproducible project foundation.", [
        "Create package folders: app, tests, scripts, data, outputs, and docs.",
        "Add pyproject.toml, .env.example, .gitignore, configuration management, and dependency locks.",
        "Set up CUDA-enabled PyTorch and a verification script that records GPU, driver, and package versions.",
        "Keep the existing implementation unchanged under research/from_scratch_sd."
    ], "A new developer can validate the environment and model availability with one command.")
    add_phase(doc, "Phase 1 - Product-Preserving Background Replacement (5-7 days)", "Prove the central e-commerce value proposition.", [
        "Accept product image, binary mask, prompt, negative prompt, seed, steps, and CFG scale.",
        "Validate image/mask dimensions, modes, prompt length, and mask coverage.",
        "Resize inputs to 512x512 and perform inpainting for the background.",
        "Use mask-aware denoising to retain product structure during generation.",
        "Composite the original foreground over the generated result using final = mask x original + (1-mask) x generated.",
        "Save final output, generated background, and metadata JSON."
    ], "The final foreground pixel-difference score is zero after compositing, while prompts visibly control the background.")
    add_phase(doc, "Phase 2 - Fast Sampling and Benchmarks (3-5 days)", "Quantify quality versus latency.", [
        "Add DDIM while retaining a baseline scheduler.",
        "Benchmark a fixed suite of 10 product/prompt/mask cases at 10, 15, 25, and 50 steps.",
        "Log end-to-end latency, p50/p95 latency, peak GPU memory, seed, scheduler, and output locations.",
        "Review prompt alignment with CLIP similarity and a small human-rating sheet."
    ], "A CSV and MLflow report identifies the best measured step count for the target laptop.")
    add_phase(doc, "Phase 3 - API, UI, and Reproducibility (5-7 days)", "Make the system usable by a non-technical merchant.", [
        "Implement FastAPI endpoints: POST /generate, GET /jobs/{job_id}, GET /outputs/{image_id}, and GET /health.",
        "Create a Gradio interface with upload controls, prompt controls, variant gallery, downloads, settings, and generation time.",
        "Persist model version, seed, scheduler, steps, guidance scale, prompt, latency, and source references in JSON and SQLite.",
        "Return clear input-validation and GPU-memory errors."
    ], "A user can launch the service, generate an ad creative, and reproduce its settings from saved metadata.")
    add_phase(doc, "Phase 4 - Tracking, Testing, and Packaging (4-6 days)", "Show MLOps discipline and reliable delivery.", [
        "Log parameters, metrics, model versions, and image artifacts to MLflow locally.",
        "Add unit tests for masks, compositing, preprocessing, metadata, and API validation.",
        "Add an integration smoke test using a mock or compact test path where feasible.",
        "Create Dockerfile and Docker Compose configuration with structured logs."
    ], "Docker Compose starts the service and automated tests pass consistently.")
    add_phase(doc, "Phase 5 - Optional Product LoRA (1-2 weeks)", "Adapt generation toward a specific product identity after the MVP is proven.", [
        "Curate 20-50 authorized product images with consistent captions.",
        "Train rank-8 LoRA adapters on cross-attention projections using FP16, batch size 1, gradient accumulation, and checkpointing.",
        "Version datasets, adapters, prompts, and validation outputs.",
        "Compare baseline and LoRA outputs on held-out prompts."
    ], "Measured, reproducible comparison demonstrates whether LoRA improves product-specific generation; do not claim improvement without results.")

    heading(doc, "Recommended Repository Structure")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    code = "LLMOPS/\n  research/from_scratch_sd/\n  app/{api, core, services, schemas, ui}/\n  tests/\n  scripts/{benchmark.py, verify_environment.py}\n  docs/{architecture.md, benchmark-results.md, model-card.md}\n  outputs/  (gitignored)\n  data/     (gitignored)\n  Dockerfile\n  compose.yaml\n  pyproject.toml\n  README.md"
    r = p.add_run(code)
    r.font.name = "Consolas"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    heading(doc, "Evaluation Metrics")
    add_table(doc, ["Goal", "Metric"], [
        ("Product preservation", "Foreground pixel difference after final composite = 0."),
        ("Prompt relevance", "CLIP text-image similarity plus human review."),
        ("Boundary quality", "Boundary-band artifact score plus visual comparison grid."),
        ("Performance", "p50/p95 latency, peak GPU memory, and images per minute."),
        ("Reliability", "Successful jobs divided by total jobs; validation and inference failures tracked."),
        ("Reproducibility", "Stored seed, model, scheduler, and parameters reproduce the same workflow.")
    ], [2600, 6760])

    heading(doc, "Final Demo Scenario")
    add_text(doc, "An online sneaker seller uploads a phone image and mask, enters a premium campaign prompt, and receives four downloadable ad creatives. The original sneaker pixels are retained while the environment is generated. Every output records its seed, model version, inference settings, and latency.")
    heading(doc, "Completion Checklist")
    add_bullets(doc, [
        "Product foreground remains pixel-identical in final output.",
        "Prompt-controlled backgrounds generate correctly for multiple test products.",
        "A 10/15/25/50-step benchmark report exists.",
        "FastAPI and Gradio demo work locally on CUDA.",
        "Outputs include reproducible JSON metadata and experiment tracking.",
        "Docker startup, tests, README, architecture diagram, and demo assets are complete.",
        "Only measured performance and quality metrics appear in the resume and README."
    ])
    heading(doc, "Resume Project Title")
    p = doc.add_paragraph()
    set_font(p.add_run("ProductStudio AI - Mask-Aware Latent Diffusion Platform for E-Commerce Ad Creative Generation"), size=12, bold=True, color=DARK_BLUE)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("ProductStudio AI Build Plan"), size=9, color="64748B")
    doc.core_properties.title = "ProductStudio AI Build Plan"
    doc.core_properties.subject = "E-commerce ad creative latent diffusion implementation roadmap"
    doc.core_properties.author = "Ankit"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
